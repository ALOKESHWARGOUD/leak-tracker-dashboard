import streamlit as st
import pandas as pd
import os
from datetime import datetime
from docx import Document
from docx2pdf import convert

# ======================
# LOAD DATA (SAFE)
# ======================
@st.cache_data
def load_data():
    path = "data/leaks.csv"

    if not os.path.exists(path):
        return pd.DataFrame(columns=[
            "title", "url", "platform", "date", "client", "severity"
        ])

    df = pd.read_csv(path)

    required_cols = ["title", "url", "platform", "date", "client", "severity"]

    for col in required_cols:
        if col not in df.columns:
            if col == "client":
                df[col] = "Doctutorials"
            elif col == "severity":
                df[col] = "Critical"
            elif col == "date":
                df[col] = pd.Timestamp.now()
            else:
                df[col] = None

    df["date"] = pd.to_datetime(df["date"], errors="coerce")

    return df


df = load_data()

# ======================
# SUMMARY
# ======================
def generate_summary(filtered, client):
    if filtered.empty:
        return None

    return {
        "client": client,
        "total_leaks": len(filtered),
        "severity": "Critical"
    }

# ======================
# WORD TEMPLATE → PDF
# ======================
def generate_word_report(summary, filtered):
    template_path = "DT_Template.docx"

    doc = Document(template_path)

    # Replace placeholders
    for para in doc.paragraphs:
        if "{{client}}" in para.text:
            para.text = para.text.replace("{{client}}", summary["client"])

        if "{{date}}" in para.text:
            para.text = para.text.replace("{{date}}", str(datetime.now().date()))

        if "{{total_leaks}}" in para.text:
            para.text = para.text.replace("{{total_leaks}}", str(summary["total_leaks"]))

        if "{{severity}}" in para.text:
            para.text = para.text.replace("{{severity}}", summary["severity"])

    # Insert table at end (simple version)
    table = doc.add_table(rows=1, cols=4)

    hdr = table.rows[0].cells
    hdr[0].text = "S.No"
    hdr[1].text = "Title"
    hdr[2].text = "Platform"
    hdr[3].text = "URL"

    for i, row in filtered.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(i + 1)
        cells[1].text = str(row["title"])
        cells[2].text = str(row["platform"])
        cells[3].text = str(row["url"])

    # Save Word file
    doc.save("report.docx")

    # Convert to PDF
    convert("report.docx", "report.pdf")

    return "report.pdf"

# ======================
# UI
# ======================
st.title("🔥 Leak Tracker Dashboard")

client = st.selectbox("Select Client", df["client"].unique())
report_type = st.radio("Report Type", ["Daily", "Weekly", "Monthly"])
date_range = st.date_input("Select Date Range", [])

filtered = df[df["client"] == client]

date_label = "N/A"

if len(date_range) == 2:
    start, end = date_range

    filtered = filtered[
        (filtered["date"] >= pd.to_datetime(start)) &
        (filtered["date"] <= pd.to_datetime(end))
    ]

    if report_type == "Daily":
        date_label = str(start)
    elif report_type == "Weekly":
        date_label = f"{start} to {end}"
    elif report_type == "Monthly":
        date_label = start.strftime("%B %Y")

# SHOW DATA
st.subheader("Filtered Data")
st.dataframe(filtered)

# GENERATE REPORT
if st.button("Generate Report"):
    summary = generate_summary(filtered, client)

    if summary:
        pdf_file = generate_word_report(summary, filtered)

        with open(pdf_file, "rb") as f:
            st.download_button(
                "⬇️ Download Report",
                f,
                file_name="Leak_Report.pdf"
            )
    else:
        st.warning("No data available for selected filters")